from sqlite3 import Connection

import docx
from docx import Document as Docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import Depends, APIRouter, Body, Path
from starlette.responses import Response
from typing_extensions import Annotated

from core.methods import get_connection

router = APIRouter()


@router.post("/{document_id}/order")
def create_order(
    document_id: Annotated[int, Path()],
    path: Annotated[str, Body()],
    connection: Annotated[Connection, Depends(get_connection)]
):
    cursor = connection.cursor()

    cursor.execute('SELECT * FROM documents WHERE id = ?', (document_id,))
    document_data = cursor.fetchone()

    cursor.execute('SELECT * FROM sports')
    sports_data = {v['id']: v['name'] for v in cursor.fetchall()}

    cursor.execute('SELECT * FROM document_athletes WHERE document_id = ? AND is_sports_category_granted = true AND is_doping_check_passed = false', (document_id,))
    athletes_data = cursor.fetchall()

    athletes_data.sort(key=lambda athlete: (sports_data[athlete['sport_id']], athlete['full_name']))

    cursor.execute('SELECT name FROM sports_categories WHERE id = ?', (document_data["sports_category_id"],))
    categories_data = cursor.fetchone()

    doc = Docx("../resources/Шаблон_приказа.docx")
    for paragraph in doc.paragraphs:
        if "%SportsСategory%" in paragraph.text:
            for run in paragraph.runs:
                if "%SportsСategory%" in run.text:
                    run.text = run.text.replace("%SportsСategory%", categories_data["name"])

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in('w:tblBorders')
    if tblBorders is None:
        tblBorders = docx.oxml.shared.OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = docx.oxml.shared.OxmlElement(f'w:{border_name}')
        border.set(docx.oxml.shared.qn('w:val'), 'single')
        border.set(docx.oxml.shared.qn('w:sz'), '4')
        border.set(docx.oxml.shared.qn('w:space'), '0')
        border.set(docx.oxml.shared.qn('w:color'), 'auto')
        tblBorders.append(border)

    for column in table.columns:
        column.width = Pt(100)

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Номер строки'
    header_cells[1].text = 'Фамилия Имя Отчество'
    header_cells[2].text = 'Дата рождения'
    header_cells[3].text = 'Наименование муниципального образования'
    header_cells[4].text = 'Организация'

    for i in range(0, 5):
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    big_counter = 0
    small_counter = 0
    cur_sport = -100
    for row_data in athletes_data:
        if cur_sport != row_data["sport_id"]:
            big_counter += 1
            if big_counter != 1:
                table.add_row()
            row_cells = table.add_row().cells
            row_cells[0].text = str(big_counter)
            for paragraph in row_cells[0].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cursor.execute('SELECT * FROM sports WHERE id = ?', (row_data["sport_id"],))
            categories_data = cursor.fetchone()
            row_cells[1].text = categories_data["name"]
            row_cells[1].merge(row_cells[4])
            for paragraph in row_cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cur_sport = row_data["sport_id"]
            small_counter = 1

        row_cells = table.add_row().cells
        row_cells[0].text = f"{big_counter}.{small_counter}"
        row_cells[1].text = row_data["full_name"]
        row_cells[2].text = row_data["birth_date"]
        row_cells[3].text = row_data["municipality"]
        row_cells[4].text = row_data["organization"]
        for i in range(0, 5):
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        small_counter += 1

    doc.save(path)

    return Response()
